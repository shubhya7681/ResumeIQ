"""
ResumeIQ - Resume Text Extraction Module
Extracts and cleans text from PDF and DOCX resume files.
"""
import os
import re


def extract_text_from_pdf(filepath):
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""


def extract_text_from_docx(filepath):
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return ""


def clean_text(text):
    """Clean and normalize extracted resume text."""
    if not text:
        return ""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,;:!?@#$%&*()\-/\\+\'\"]+', '', text)
    # Remove extra newlines
    text = re.sub(r'\n\s*\n', '\n', text)
    return text.strip()


def extract_text(filepath):
    """
    Main extraction function. Detects file type and extracts text accordingly.
    Returns cleaned text ready for analysis.
    """
    if not os.path.exists(filepath):
        return ""

    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        raw_text = extract_text_from_pdf(filepath)
    elif ext == '.docx':
        raw_text = extract_text_from_docx(filepath)
    else:
        return ""

    return clean_text(raw_text)


def extract_sections(text):
    """
    Attempt to identify common resume sections from the text.
    Returns a dictionary of section names and their content.
    """
    sections = {
        'objective': '',
        'summary': '',
        'education': '',
        'experience': '',
        'skills': '',
        'projects': '',
        'certifications': '',
        'achievements': '',
        'contact': '',
        'references': ''
    }

    # Common section header patterns
    section_patterns = {
        'objective': r'(?i)(career\s*objective|objective|career\s*goal)',
        'summary': r'(?i)(professional\s*summary|summary|profile|about\s*me)',
        'education': r'(?i)(education|academic|qualification|degree)',
        'experience': r'(?i)(experience|work\s*history|employment|professional\s*experience|work\s*experience)',
        'skills': r'(?i)(skills|technical\s*skills|competencies|technologies|tools)',
        'projects': r'(?i)(projects|academic\s*projects|personal\s*projects)',
        'certifications': r'(?i)(certifications?|certificates?|licenses?|courses?)',
        'achievements': r'(?i)(achievements?|awards?|honors?|accomplishments?)',
        'contact': r'(?i)(contact|personal\s*information|personal\s*details)',
        'references': r'(?i)(references?)'
    }

    text_lower = text.lower()
    found_positions = []

    for section_name, pattern in section_patterns.items():
        match = re.search(pattern, text)
        if match:
            found_positions.append((match.start(), section_name))

    # Sort by position in text
    found_positions.sort(key=lambda x: x[0])

    # Extract content between section headers
    for i, (pos, section_name) in enumerate(found_positions):
        if i + 1 < len(found_positions):
            next_pos = found_positions[i + 1][0]
            sections[section_name] = text[pos:next_pos].strip()
        else:
            sections[section_name] = text[pos:].strip()

    return sections
